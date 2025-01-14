from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from utils import get_numberstring_three_dig, get_numberstring_two_dig
import os

#functions
# Función para agregar bordes a una celda
def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        start={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        end={"sz": 12, "val": "single", "color": "000000", "space": "0"},
    )
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for border_name in ["start", "top", "end", "bottom"]:
        if border_name in kwargs:
            border = OxmlElement(f"w:{border_name}")
            for attr_name, attr_value in kwargs[border_name].items():
                border.set(qn(f"w:{attr_name}"), str(attr_value))
            tcPr.append(border)

# Función para agregar color de fondo a una celda
def set_cell_background(cell, color):
    """
    Set cell background color
    Usage: set_cell_background(cell, "D9D9D9")
    """
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)  # Color de fondo
    tc_pr.append(shd)

def format_cell(cell, font_size=9, bold=False):
    """
    Aplica formato a la celda: alineación vertical, tamaño de fuente y opcionalmente negrita.
    """
    # Centrar verticalmente el contenido de la celda
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Cambiar el tamaño de la fuente y aplicar estilo
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.bold = bold



def create_word_document(serie, name):  
  # Crear el documento
  doc = Document()

  section = doc.sections[0]
  section.left_margin = Inches(0.5) 
  section.right_margin = Inches(0.5)
  section.top_margin = Inches(0)  
  section.bottom_margin = Inches(0.5)

  # Estilos generales para el documento
  style = doc.styles['Normal']
  font = style.font
  font.name = 'Arial'
  font.size = Pt(10)

  # Agregar un encabezado
  header = doc.sections[0].header
  header.height = Inches(0) 

  # Agregar una tabla para el encabezado (1 fila, 2 columnas)
  header_table = header.add_table(rows=1, cols=2, width=Inches(7))
  header_table.autofit = False
  header_table.columns[0].width = Inches(2)  # Ajustar el ancho de la columna para el logo
  header_table.columns[1].width = Inches(6.5)  # Ajustar el ancho de la columna para el texto

  # Agregar el texto del encabezado
  text_cell = header_table.cell(0, 0)
  text_paragraph = text_cell.paragraphs[0]
  text_paragraph.add_run("Vicepresidencia Tecnología e Infraestructura\n"
                        "Gerencia General Proyectos Mayores\n"
                        "Coordinación de Cicre. Región Capital")
  text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinear el texto a la derecha

  # Agregar el logo (asegúrate de que la ruta sea correcta)
  logo_cell = header_table.cell(0, 1)
  logo_paragraph = logo_cell.paragraphs[0]
  logo_run = logo_paragraph.add_run()
  logo_run.add_picture('./static/images/logoCantv.png', width=Inches(1.41), height=Inches(0.69))  # Ajusta el tamaño del logo
  logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Alinear el logo a la izquierda


  # Título
  title = doc.add_paragraph()
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER

  # Agregar texto subrayado
  run1 = title.add_run("ACTA DE ACEPTACIÓN\n")
  run1.bold = True  # Si deseas que el texto sea negrita
  run1.underline = True  # Agregar subrayado

  run2 = title.add_run("TERMINAL DE ACCESO DE FIBRA (FAT)")
  run2.bold = True  # Si deseas que el texto sea negrita
  run2.underline = True  # Agregar subrayado

  #============================================================================
  # Crear tabla principal con datos
  table = doc.add_table(rows=9, cols=5)
  table.autofit = False
  
  for cell in table.columns[0].cells:
    cell.width = Inches(1.8)
    
  # Añadir contenido a las celdas (ejemplo)
  table.cell(0, 0).text = "Contratista/Proveedor"
  table.cell(0, 1).merge(table.cell(0, 4))
  table.cell(0, 1).text = "TELECOMUNICACIONES SENATEL"
  format_cell(table.cell(0, 0)) 
  format_cell(table.cell(0, 1))

  table.cell(1, 0).text = "N° Elemento PEP"
  table.cell(1, 1).merge(table.cell(1, 4))
  table.cell(1, 1).text = "RED/24-94-007-P-FI001"
  format_cell(table.cell(1, 0))  # Aplicar formato
  format_cell(table.cell(1, 1))

  table.cell(2, 0).text = "N° Orden de Pedido"
  table.cell(2, 1).merge(table.cell(2, 4))
  table.cell(2, 1).text = ""
  format_cell(table.cell(2, 0))
  format_cell(table.cell(2, 1))

  table.cell(3, 0).text = "FAT N° Identificación"
  table.cell(3, 1).text = ""
  table.cell(3, 2).merge(table.cell(3, 3))
  table.cell(3, 2).text = "Cantidad de Puertos a ser certificados"
  table.cell(3, 4).text = ""
  format_cell(table.cell(3, 0))
  format_cell(table.cell(3, 1))
  format_cell(table.cell(3, 2))
  format_cell(table.cell(3, 4))

  table.cell(4, 0).text = "Localidad"
  table.cell(4, 1).merge(table.cell(4, 4)) 
  table.cell(4, 1).text = str(serie['sectores'])
  format_cell(table.cell(4, 0))
  format_cell(table.cell(4, 1))

  table.cell(5, 0).text = "Dirección"
  table.cell(5, 1).merge(table.cell(5, 4)) 
  table.cell(5, 1).text = str(serie['direccion_inm'])
  format_cell(table.cell(5, 0))
  format_cell(table.cell(5, 1))

  table.cell(6, 0).text = "Coordenadas Geográficas"
  table.cell(6, 1).text = "Latitud:"
  table.cell(6, 2).text = str(serie['y'])
  table.cell(6, 3).text = "Longitud:"
  table.cell(6, 4).text = str(serie['X'])
  format_cell(table.cell(6, 0))
  format_cell(table.cell(6, 1))
  format_cell(table.cell(6, 2))
  format_cell(table.cell(6, 3))

  table.cell(7, 0).text = "Nombre del Proyecto y Breve Descripción de la Obra"
  table.cell(7, 1).merge(table.cell(7, 3)) 
  table.cell(7, 1).text = ""
  format_cell(table.cell(7, 0))
  format_cell(table.cell(7, 1))

  table.cell(8, 0).text = "Unidad que recibe o Certifica"
  table.cell(8, 1).merge(table.cell(8, 3)) 
  table.cell(8, 1).text = ""
  format_cell(table.cell(8, 0))
  format_cell(table.cell(8, 1))

  ## estilando tabla
  for row in table.rows:
      for cell in row.cells:
          set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                              bottom={"sz": 12, "val": "single", "color": "000000"},
                              start={"sz": 12, "val": "single", "color": "000000"},
                              end={"sz": 12, "val": "single", "color": "000000"})

  for i in range(0,9):
    set_cell_background(table.cell(i, 0), "D9D9D9")

  set_cell_background(table.cell(3, 2), "D9D9D9")
  set_cell_background(table.cell(6, 1), "D9D9D9")
  set_cell_background(table.cell(6, 3), "D9D9D9")

  table.alignment = WD_TABLE_ALIGNMENT.CENTER

  # Añadir otro contenido, como observaciones o descripciones
  doc.add_paragraph("\nEn la ciudad de Caracas,  a  los  días del mes de XXXX del año 2023, se constituyeron los representantes abajo firmantes, a los fines de"
                    "dejar constancia de la Aceptación Provisional descrita en el epígrafe.")

  first_item = doc.add_paragraph()
  run = first_item.add_run("• PRIMERO: ")
  run.bold = True
  run.underline = True
  first_item.add_run(
      "Se efectuaron las pruebas de aceptación correspondientes y arrojaron "
      "los resultados detallados en el Protocolo de Prueba correspondientes, quedando sin reparos "
      "(Los resultados de las pruebas se encuentran detallados en la hoja."
  )
  first_item.paragraph_format.left_indent = Pt(36)

  secnd_item = doc.add_paragraph()
  run = secnd_item.add_run("• SEGUNDO: ")
  run.bold = True
  run.underline = True
  secnd_item.add_run(
      "Asimismo, se deja constancia de la ejecución de los trabajos aplicando la normativa "
      "de calidad establecida por CANTV.\n"
  )
  secnd_item.paragraph_format.left_indent = Pt(36)



  # Añadir otra tabla para las firmas
  signature_table = doc.add_table(rows=3, cols=3)

  # Cabecera de firmas
  signature_table.cell(0, 0).text = "En representación de CANTV: Por Operaciones Regionales"
  signature_table.cell(0, 1).text = "En representación de CANTV: Unidad: Coordinador Cicre Capital"
  signature_table.cell(0, 2).text = "En representación de Contratista / Proveedor"

  # Añadir campos para firmas
  signature_table.cell(1, 0).text = f"Nombre:\nN° SAP:\nC.I.:\nFirma: __________________"
  signature_table.cell(1, 1).text = "Nombre:\nN° SAP:\nC.I.:\nFirma: __________________"
  signature_table.cell(1, 2).text = "Contratista:\nNombre:\nC.I.:\nFirma: __________________"

  # Bordes exteriores: Primera y última fila (bordes superior e inferior)
  for col in range(3):
      # Bordes para la primera fila (superior)
      set_cell_border(signature_table.cell(0, col), top={"sz": 12, "val": "single", "color": "000000"})
      # Bordes para la última fila (inferior)
      set_cell_border(signature_table.cell(2, col), bottom={"sz": 12, "val": "single", "color": "000000"})

  # Bordes exteriores: Primera y última columna (bordes izquierdo y derecho)
  for row in range(3):
      # Bordes para la primera columna (izquierda)
      set_cell_border(signature_table.cell(row, 0), start={"sz": 12, "val": "single", "color": "000000"})
      # Bordes para la última columna (derecha)
      set_cell_border(signature_table.cell(row, 2), end={"sz": 12, "val": "single", "color": "000000"})

  for row in range(3):
      for col in range(2):  # Aplica solo a las columnas 0 y 1
          set_cell_border(signature_table.cell(row, col), end={"sz": 12, "val": "single", "color": "000000"})
          
  signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER

  doc.add_paragraph()

  # Añadir otra tabla
  end_table = doc.add_table(rows=1, cols=4)
  end_table.autofit = False
  
  for cell in end_table.columns[0].cells:
    cell.width = Inches(2.1)
    
  # Cabecera de firmas
  end_table.cell(0, 0).text = "Retiraron escombros de la obra"
  end_table.cell(0, 1).text = "Si "
  end_table.cell(0, 2).text = "No "
  end_table.cell(0, 3).text = "Observaciones: "

  end_table.columns[0].width = Inches(2.5)
  end_table.columns[1].width = Inches(1)
  end_table.columns[2].width = Inches(1)
  end_table.columns[3].width = Inches(2.5)

  for row in range(1):
      for col in range(4):
          # Bordes exteriores: celdas en los extremos de la tabla
          if col == 0:  # Primera columna (izquierda)
              set_cell_border(end_table.cell(row, col),
                              top={"sz": 12, "val": "single", "color": "000000"},
                              bottom={"sz": 12, "val": "single", "color": "000000"},
                              start={"sz": 12, "val": "single", "color": "000000"})
          elif col == 3:  # Última columna (derecha)
              set_cell_border(end_table.cell(row, col),
                              top={"sz": 12, "val": "single", "color": "000000"},
                              bottom={"sz": 12, "val": "single", "color": "000000"},
                              end={"sz": 12, "val": "single", "color": "000000"})
          else:  # Otras celdas (solo bordes superior e inferior)
              set_cell_border(end_table.cell(row, col),
                              top={"sz": 12, "val": "single", "color": "000000"},
                              bottom={"sz": 12, "val": "single", "color": "000000"})
      
      # Agregar borde derecho adicional a la primera celda (0,0)
      set_cell_border(end_table.cell(0, 0), 
                      end={"sz": 12, "val": "single", "color": "000000"})
      

  end_table.alignment = WD_TABLE_ALIGNMENT.CENTER

  #=====================================================================================================================================================
  # hoja2
  #=====================================================================================================================================================

  doc.add_page_break()

  # Tabla 1 (Puerto, Potencia, etc.)
  # Crear tabla principal con datos
  table1 = doc.add_table(rows=3, cols=4)
  table1.autofit = False
  
  for cell in table1.columns[0].cells:
    cell.width = Inches(0.7)
  for cell in table1.columns[2].cells:
    cell.width = Inches(0.7)

  # Añadir contenido a las celdas (ejemplo)
  table1.cell(0, 0).merge(table1.cell(0, 3))
  table1.cell(0, 0).text = "FAT 1  Puertos de salida del Splitter"
  format_cell(table1.cell(0, 0))


  # Fila 2
  table1.cell(1, 0).text = "Puerto"
  table1.cell(1, 1).text = "Potencia Recibida (dBm)"
  table1.cell(1, 2).text = "Puerto"
  table1.cell(1, 3).text = "Potencia Recibida (dBm)"
  format_cell(table1.cell(1, 0)) 
  format_cell(table1.cell(1, 1))
  format_cell(table1.cell(1, 2))
  format_cell(table1.cell(1, 3))

  # Fila 3
  table1.cell(2, 0).text = ""
  table1.cell(2, 1).text = ""
  table1.cell(2, 2).text = ""
  table1.cell(2, 3).text = ""
  format_cell(table1.cell(2, 0)) 
  format_cell(table1.cell(2, 1))
  format_cell(table1.cell(2, 2))
  format_cell(table1.cell(2, 3))
  
  for row in table1.rows:
    for cell in row.cells:
      set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                      bottom={"sz": 12, "val": "single", "color": "000000"},
                      start={"sz": 12, "val": "single", "color": "000000"},
                      end={"sz": 12, "val": "single", "color": "000000"})

  # Aplicar fondo gris a las filas 0 y 1
  for cell in table1.row_cells(0):
      set_cell_background(cell, "D9D9D9")  # Color gris

  for cell in table1.row_cells(1):
      set_cell_background(cell, "D9D9D9")  # Color gris

  # Aplicar fondo gris a las columnas 0 y 2
  for row in table1.rows:
      set_cell_background(row.cells[0], "D9D9D9")  # Columna 0
      set_cell_background(row.cells[2], "D9D9D9")  # Columna 2

  doc.add_paragraph()
  #=========================================== Tabla 2 (Código Catastral)
  # Crear tabla principal con datos
  table2 = doc.add_table(rows=5, cols=6)
  table2.autofit = False
  
  for cell in table2.columns[0].cells:
    cell.width = Inches(1.5)
  for cell in table2.columns[2].cells:
    cell.width = Inches(1.5)
  for cell in table2.columns[4].cells:
    cell.width = Inches(1.5)
  for cell in table2.columns[1].cells:
    cell.width = Inches(0.8)
  for cell in table2.columns[3].cells:
    cell.width = Inches(0.8)
  for cell in table2.columns[5].cells:
    cell.width = Inches(0.8)

  #Fila 1
  # Añadir contenido a las celdas (ejemplo)
  table2.cell(0, 0).merge(table2.cell(0, 5))
  table2.cell(0, 0).text = "Código Catastral "
  format_cell(table2.cell(0, 0))


  # Fila 2
  table2.cell(1, 0).text = "Estado"
  table2.cell(1, 1).text = get_numberstring_two_dig(serie['cod_estado'])
  table2.cell(1, 2).text = "Ciudad"
  table2.cell(1, 3).text = ""
  table2.cell(1, 4).text = "Municipio"
  table2.cell(1, 5).text = get_numberstring_two_dig(serie['cod_municipio'])
  format_cell(table2.cell(1, 0)) 
  format_cell(table2.cell(1, 1))
  format_cell(table2.cell(1, 2))
  format_cell(table2.cell(1, 3))
  format_cell(table2.cell(1, 4))
  format_cell(table2.cell(1, 5))

  # Fila 3
  table2.cell(2, 0).text = "Parroquia"
  table2.cell(2, 1).text = get_numberstring_two_dig(serie['cod_parroquia'])
  table2.cell(2, 2).text = "Urbanización"
  table2.cell(2, 3).text = ""
  table2.cell(2, 4).text = "Ámbito"
  table2.cell(2, 5).text = get_numberstring_three_dig(serie['cod_ambito'])
  format_cell(table2.cell(2, 0)) 
  format_cell(table2.cell(2, 1))
  format_cell(table2.cell(2, 2))
  format_cell(table2.cell(2, 3))
  format_cell(table2.cell(2, 4))
  format_cell(table2.cell(2, 5))


  # Fila 4
  table2.cell(3, 0).text = "Sector"
  table2.cell(3, 1).text = get_numberstring_three_dig(serie['cod_sector'])
  table2.cell(3, 2).text = "Manzana"
  table2.cell(3, 3).text = get_numberstring_three_dig(serie['cod_manz'])
  table2.cell(3, 4).text = "Parcela"
  table2.cell(3, 5).text = get_numberstring_three_dig(serie['cod_parc'])
  format_cell(table2.cell(3, 0)) 
  format_cell(table2.cell(3, 1))
  format_cell(table2.cell(3, 2))
  format_cell(table2.cell(3, 3))
  format_cell(table2.cell(3, 4))
  format_cell(table2.cell(3, 5))

  # Fila 5
  table2.cell(4, 0).text = "Sub-Parcela"
  table2.cell(4, 1).text = get_numberstring_three_dig(serie['cod_campo_edif'])
  table2.cell(4, 2).text = "Código Postal"
  table2.cell(4, 3).text = ""
  table2.cell(4, 4).text = ""
  table2.cell(4, 5).text = ""
  format_cell(table2.cell(4, 0)) 
  format_cell(table2.cell(4, 1))
  format_cell(table2.cell(4, 2))
  format_cell(table2.cell(4, 3))
  format_cell(table2.cell(4, 4))
  format_cell(table2.cell(4, 5))
  
  for row in table2.rows:
    for cell in row.cells:
      set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                      bottom={"sz": 12, "val": "single", "color": "000000"},
                      start={"sz": 12, "val": "single", "color": "000000"},
                      end={"sz": 12, "val": "single", "color": "000000"})
  
  for cell in table2.row_cells(0):
    set_cell_background(cell, "D9D9D9")  # Color gris para la fila 1

# Sombrear las columnas 0, 2, y 4 en todas las filas
  for row in table2.rows:
    set_cell_background(row.cells[0], "D9D9D9")  # Columna 0
    set_cell_background(row.cells[2], "D9D9D9")  # Columna 2
    set_cell_background(row.cells[4], "D9D9D9")  # Columna 4

  doc.add_paragraph()
  #=========================================== Tabla 3 (Código Catastral)

  # Crear tabla principal con datos
  table3 = doc.add_table(rows=2, cols=3)
  table3.autofit = False
  
  for cell in table3.columns[0].cells:
    cell.width = Inches(0.8)
    
  #Fila 1
  # Añadir contenido a las celdas (ejemplo)
  table3.cell(0, 0).merge(table3.cell(0, 2))
  table3.cell(0, 0).text = "Ubicación"
  format_cell(table3.cell(0, 0))


  # Fila 2
  table3.cell(1, 0).text = "Orientación"
  table3.cell(1, 1).text = ""
  table3.cell(1, 2).text = ""
  format_cell(table3.cell(1, 0)) 
  format_cell(table3.cell(1, 1))
  format_cell(table3.cell(1, 2))
  
  set_cell_background(table3.cell(1, 0), "D9D9D9")

  # Aplicar bordes a todas las celdas en table3
  for row in table3.rows:
    for cell in row.cells:
      set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                      bottom={"sz": 12, "val": "single", "color": "000000"},
                      start={"sz": 12, "val": "single", "color": "000000"},
                      end={"sz": 12, "val": "single", "color": "000000"})

  doc.add_paragraph()
  #=========================================== Tabla 4 (Código Catastral)

  # Crear tabla principal con datos
  table4 = doc.add_table(rows=2, cols=6)
  table4.autofit = False
  
  for cell in table4.columns[0].cells:
    cell.width = Inches(0.7)
  for cell in table4.columns[2].cells:
    cell.width = Inches(0.7)
  for cell in table4.columns[4].cells:
    cell.width = Inches(0.7)

  #Fila 1
  # Añadir contenido a las celdas (ejemplo)
  table4.cell(0, 0).merge(table4.cell(0, 5))
  table4.cell(0, 0).text = "Uso Urbano"
  format_cell(table4.cell(0, 0))


  # Fila 2
  table4.cell(1, 0).text = "Uso Urbano "
  table4.cell(1, 1).text = str(serie['uso_inmueble'])
  table4.cell(1, 2).text = "Tipo Inmueble"
  table4.cell(1, 3).text = str(serie['tipo_edificacion'])
  table4.cell(1, 4).text = "Nombre / Razón Social"
  table4.cell(1, 5).text = str(serie['nombreedificacion'])
  format_cell(table4.cell(1, 0)) 
  format_cell(table4.cell(1, 1))
  format_cell(table4.cell(1, 2))
  format_cell(table4.cell(1, 3))
  format_cell(table4.cell(1, 4))
  format_cell(table4.cell(1, 5))
  
  for col in range(6):
    set_cell_background(table4.cell(1, col), "D9D9D9")

  # Sombrear columnas 0, 2 y 4 en table4, fila 1
  set_cell_background(table4.cell(1, 0), "D9D9D9")
  set_cell_background(table4.cell(1, 2), "D9D9D9")
  set_cell_background(table4.cell(1, 4), "D9D9D9")

  # Aplicar bordes a todas las celdas en table4
  for row in table4.rows:
    for cell in row.cells:
      set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                      bottom={"sz": 12, "val": "single", "color": "000000"},
                      start={"sz": 12, "val": "single", "color": "000000"},
                      end={"sz": 12, "val": "single", "color": "000000"})


  doc.add_paragraph()
  #=========================================== Tabla 5 (Código Catastral)

  # Crear tabla principal con datos
  table5 = doc.add_table(rows=5, cols=4)
  table5.autofit = False
  
  for cell in table5.columns[0].cells:
    cell.width = Inches(0.4)
  for cell in table5.columns[2].cells:
    cell.width = Inches(0.4)

  #Fila 1
  # Añadir contenido a las celdas (ejemplo)
  table5.cell(0, 0).merge(table5.cell(0, 3))
  table5.cell(0, 0).text = "RESIDENCIAS UNIFAMILIARES  ATENDIDAS Y JURÍDICO"
  format_cell(table5.cell(0, 0))


  # Fila 2
  table5.cell(1, 0).text = "Nº"
  table5.cell(1, 1).text = "Nombre"
  table5.cell(1, 2).text = "Nº"
  table5.cell(1, 3).text = "Nombre"
  format_cell(table5.cell(1, 0)) 
  format_cell(table5.cell(1, 1))
  format_cell(table5.cell(1, 2))
  format_cell(table5.cell(1, 3))

  table5.cell(2, 0).text = "1"
  table5.cell(2, 1).text = ""
  table5.cell(2, 2).text = "4"
  table5.cell(2, 3).text = ""
  format_cell(table5.cell(2, 0)) 
  format_cell(table5.cell(2, 1))
  format_cell(table5.cell(2, 2))
  format_cell(table5.cell(2, 3))

  table5.cell(3, 0).text = "2"
  table5.cell(3, 1).text = ""
  table5.cell(3, 2).text = "5"
  table5.cell(3, 3).text = ""
  format_cell(table5.cell(3, 0)) 
  format_cell(table5.cell(3, 1))
  format_cell(table5.cell(3, 2))
  format_cell(table5.cell(3, 3))

  table5.cell(4, 0).text = "3"
  table5.cell(4, 1).text = ""
  table5.cell(4, 2).text = "6"
  table5.cell(4, 3).text = ""
  format_cell(table5.cell(4, 0)) 
  format_cell(table5.cell(4, 1))
  format_cell(table5.cell(4, 2))
  format_cell(table5.cell(4, 3))

  for col in range(4):
    set_cell_background(table5.cell(0, col), "D9D9D9")

  # Sombrear las columnas 0 y 2 para todas las filas (excepto la primera fila que ya fue sombreada)
  for row in range(1, 5):
      set_cell_background(table5.cell(row, 0), "D9D9D9")
      set_cell_background(table5.cell(row, 2), "D9D9D9")

  # Aplicar bordes a todas las celdas de la tabla
  for row in table5.rows:
    for cell in row.cells:
      set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                      bottom={"sz": 12, "val": "single", "color": "000000"},
                      start={"sz": 12, "val": "single", "color": "000000"},
                      end={"sz": 12, "val": "single", "color": "000000"})
 
  #=====================================================================================================================================================
  # hoja3
  #=====================================================================================================================================================

  doc.add_page_break()

  # Tabla 1 (Puerto, Potencia, etc.)
  # Crear tabla principal con datos
  table_page3 = doc.add_table(rows=18, cols=4)
  table_page3.autofit = False

  # Añadir contenido a las celdas (ejemplo)
  table_page3.cell(0, 0).merge(table_page3.cell(0, 3))
  table_page3.cell(0, 0).text = "FAT 72.1  Puertos de salida del Splitter"
  format_cell(table_page3.cell(0, 0))


  # Fila 2
  table_page3.cell(1, 0).text = "Puerto"
  table_page3.cell(1, 1).text = "Potencia Recibida (dBm)"
  table_page3.cell(1, 2).text = "Puerto"
  table_page3.cell(1, 3).text = "Potencia Recibida (dBm)"
  format_cell(table_page3.cell(1, 0)) 
  format_cell(table_page3.cell(1, 1))
  format_cell(table_page3.cell(1, 2))
  format_cell(table_page3.cell(1, 3))

  for i in range(16):
    table_page3.cell(i + 2, 0).text = str(i + 2)
    table_page3.cell(i + 2, 1).text = ""
    table_page3.cell(i + 2, 2).text = str(17 + i)
    table_page3.cell(i + 2, 3).text = ""
    format_cell(table_page3.cell(i + 2, 0)) 
    format_cell(table_page3.cell(i + 2, 1))
    format_cell(table_page3.cell(i + 2, 2))
    format_cell(table_page3.cell(i + 2, 3))

  table_page3.columns[0].width = Inches(0.5)
  table_page3.columns[1].width = Inches(1)
  table_page3.columns[2].width = Inches(0.5)
  table_page3.columns[3].width = Inches(1)
  
  for row in table_page3.rows:
    for cell in row.cells:
      set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        start={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        end={"sz": 12, "val": "single", "color": "000000", "space": "0"},
      )

  # Sombrear las filas 0 y 1
  for j in range(2):  # Filas 0 y 1
    for i in range(4):  # Todas las columnas de la fila
      set_cell_background(table_page3.cell(j, i), "D9D9D9")  # Color gris claro

  # Sombrear las columnas 0 y 2 en todas las filas
  for i in range(18):  # Todas las filas
    set_cell_background(table_page3.cell(i, 0), "D9D9D9")  # Columna 0
    set_cell_background(table_page3.cell(i, 2), "D9D9D9")  # Columna 2
    
  name = 'Acta_Aceptacion_' + name.replace(" ", "_") + '.docx'
  base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'data', 'actaAceptacionWord'))
  output_path = os.path.join(base_dir, name)
  doc.save(output_path)